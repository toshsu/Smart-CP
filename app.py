# app.py
import io
from typing import List
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import StreamingResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
from docx.shared import Pt
from nlp import parse_key_values, parse_clauses, merge_clauses, detect_gaps, fill_placeholders, normalize

app = FastAPI(title="Smart CP Generator", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/api/health")
def health():
    return {"status": "ok"}

@app.post("/api/generate")
async def generate_cp(
    recap: UploadFile = File(..., description="Fixture Recap (.txt)"),
    base_cp: UploadFile = File(..., description="Base CP Template (.docx)"),
    negotiated: UploadFile = File(..., description="Negotiated Clauses (.txt)"),
    filename: str = Form("Final_CP.docx")
):
    recap_txt = normalize((await recap.read()).decode("utf-8", errors="ignore"))
    negotiated_txt = normalize((await negotiated.read()).decode("utf-8", errors="ignore"))
    base_bytes = await base_cp.read()

    # Extract data
    recap_kv = parse_key_values(recap_txt)
    negotiated_kv = parse_key_values(negotiated_txt)
    # negotiated key-values override recap ones
    merged_kv = {**recap_kv, **negotiated_kv}

    # Clauses
    base_doc = Document(io.BytesIO(base_bytes))
    base_text = "\n".join(p.text for p in base_doc.paragraphs)
    base_clauses = parse_clauses(base_text)

    negotiated_clauses = parse_clauses(negotiated_txt)

    merged_clauses, conflicts = merge_clauses(base_clauses, negotiated_clauses)
    gaps = detect_gaps(merged_clauses)

    # Build final Document
    final_doc = Document(io.BytesIO(base_bytes))
    # Replace placeholders everywhere
    for paragraph in final_doc.paragraphs:
        if "{{" in paragraph.text:
            for run in paragraph.runs:
                run.text = fill_placeholders(run.text, merged_kv)

    # Rebuild a "Clauses" section at the end for clarity
    final_doc.add_page_break()
    h = final_doc.add_paragraph("Clauses").style = final_doc.styles['Heading 1']
    for cid, body in merged_clauses:
        p = final_doc.add_paragraph()
        run = p.add_run(f"{cid}. {body}")
        run.font.size = Pt(11)

    # Build a validation report
    report = {
        "placeholders_filled": merged_kv,
        "conflicts": [{"clause_id": c, "base": b, "negotiated": n} for c, b, n in conflicts],
        "gaps": gaps,
        "total_clauses": len(merged_clauses),
    }

    # Prepare response as a ZIP (docx + json report)
    mem = io.BytesIO()
    with zipfile.ZipFile(mem, mode="w", compression=zipfile.ZIP_DEFLATED) as z:
        # Save docx
        docx_bytes = io.BytesIO()
        final_doc.save(docx_bytes)
        z.writestr(filename, docx_bytes.getvalue())
        z.writestr("validation_report.json", json.dumps(report, indent=2))
    mem.seek(0)

    return StreamingResponse(mem, media_type="application/zip", headers={
        "Content-Disposition": f'attachment; filename="cp_bundle.zip"'
    })
